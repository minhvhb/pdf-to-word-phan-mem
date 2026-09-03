import yaml
from yaml.loader import SafeLoader
import streamlit_authenticator as stauth
import streamlit as st

# --- BẮT ĐẦU ĐOẠN LẤY MÃ ---
mk = stauth.Hasher.hash_passwords(['Cns@2026@'])
st.error(f"Hãy Copy đoạn mã này: {mk[0]}")
st.stop()
# --- KẾT THÚC ĐOẠN LẤY MÃ ---

# 1. Đọc dữ liệu tài khoản từ file config
with open('config.yaml') as file:
    config = yaml.load(file, Loader=SafeLoader)

# 2. Cài đặt công cụ đăng nhập (đã xóa pre-authorized theo bản cập nhật mới)
authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days']
)

# 3. Tạo khung đăng nhập
authenticator.login()

# 4. Kiểm tra trạng thái đăng nhập
if st.session_state["authentication_status"] == False:
    st.error('Tên đăng nhập hoặc mật khẩu không đúng!')
    st.stop()
elif st.session_state["authentication_status"] == None:
    st.warning('Vui lòng đăng nhập để sử dụng công cụ')
    st.stop()

# Đăng nhập thành công
authenticator.logout('Đăng xuất', 'main')
st.write(f'Chào mừng **{st.session_state["name"]}**!')
st.divider()

# --- TỪ DÒNG NÀY TRỞ XUỐNG LÀ CODE APP PDF CŨ CỦA BẠN (GIỮ NGUYÊN) ---


import streamlit as st
import os
import re
import pandas as pd
from io import BytesIO
import json
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.cell.rich_text import CellRichText, TextBlock
from openpyxl.cell.text import InlineFont
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from PIL import Image
from pypdf import PdfReader, PdfWriter, Transformation
import difflib

# ==========================================
# 1. CẤU HÌNH TRANG CHỦ ĐẠO
# ==========================================
st.set_page_config(page_title="Hệ thống Công cụ AI", page_icon="⚙️", layout="centered", initial_sidebar_state="expanded")
st.markdown(
    """
    <style>
        [data-testid="collapsedControl"] { display: none; }
        [data-testid="stSidebarCollapseButton"] { display: none; }
    </style>
    """,
    unsafe_allow_html=True
)

if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0

def clear_file():
    st.session_state.uploader_key += 1

# ==========================================
# 2. KHU VỰC APP 1: CHUYỂN PDF SANG WORD
# ==========================================
def parse_and_add_runs(paragraph, text):
    parts_bold = re.split(r'\*\*(.*?)\*\*', text)
    for i, p_bold in enumerate(parts_bold):
        is_bold = (i % 2 == 1)
        parts_italic = re.split(r'\*(.*?)\*', p_bold)
        for j, p_italic in enumerate(parts_italic):
            is_italic = (j % 2 == 1)
            parts_underline = re.split(r'<u>(.*?)</u>', p_italic)
            for k, p_underline in enumerate(parts_underline):
                is_underline = (k % 2 == 1)
                if p_underline:
                    run = paragraph.add_run(p_underline)
                    run.bold = is_bold
                    run.italic = is_italic
                    run.underline = is_underline

def clean_tags_and_align(text, paragraph, default_align):
    align = default_align
    text_clean = text.strip()
    
    if '[CENTER]' in text_clean:
        align = WD_ALIGN_PARAGRAPH.CENTER
        text_clean = text_clean.replace('[CENTER]', '').strip()
    if '[RIGHT]' in text_clean:
        align = WD_ALIGN_PARAGRAPH.RIGHT
        text_clean = text_clean.replace('[RIGHT]', '').strip()
    if '[LEFT]' in text_clean:
        align = WD_ALIGN_PARAGRAPH.LEFT
        text_clean = text_clean.replace('[LEFT]', '').strip()
        
    paragraph.alignment = align
    return text_clean

def app_pdf_to_word():
    try:
        api_key_input = st.secrets["GEMINI_API_KEY"]
    except KeyError:
        st.error("⚠️ Hệ thống chưa được cấu hình API Key. Vui lòng liên hệ Quản trị viên!")
        st.stop()

    st.title("📄 Ứng dụng Chuyển đổi PDF & Ảnh sang Word")
    st.markdown("Sử dụng **Google Gemini AI** để trích xuất văn bản và bảng biểu chuẩn Form Hành chính.")

    uploaded_file = st.file_uploader("Tải lên file ảnh (JPG, PNG) hoặc PDF:", type=["jpg", "jpeg", "png", "pdf"], key=f"app1_{st.session_state.uploader_key}")

    if uploaded_file is not None:
        st.success(f"Đã tải lên file: **{uploaded_file.name}**")
        
        if st.button("🚀 Bắt đầu Chuyển đổi", type="primary"):
            with st.spinner("🤖 AI đang phân tích lề và vẽ lại trang Word khổ A4..."):
                try:
                    temp_input_path = f"temp_{uploaded_file.name}"
                    with open(temp_input_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())

                    client = genai.Client(api_key=api_key_input)
                    with open(temp_input_path, "rb") as f:
                        file_bytes = f.read()
                    mime_type = "application/pdf" if uploaded_file.name.endswith(".pdf") else "image/jpeg"
                    
                    prompt = """
                    NHIỆM VỤ OCR - BẮT BUỘC TUÂN THỦ NGHIÊM NGẶT CÁC QUY TẮC SAU:

                    1. THỂ THỨC VĂN BẢN (QUỐC HIỆU & CHỮ KÝ):
                       - Phần Quốc hiệu trên cùng (hoặc các khối chữ ký nằm song song ngang nhau): BẮT BUỘC dùng [HEADER_TABLE] (bảng tàng hình 2 cột) để chia tỷ lệ.
                       - TUYỆT ĐỐI KHÔNG dùng thẻ HTML (như <td>, <colspan>...).
                       - TUYỆT ĐỐI KHÔNG vẽ bảng cho phần Footer (ví dụ: BM/66-VPCQ, Lần ban hành, Trang). Hãy viết thành văn bản thường trên 1 dòng.

                    2. CANH LỀ ĐOẠN VĂN (Sử dụng Tag):
                       - Ghi [CENTER] ở đầu MỌI dòng cần canh giữa.
                       - Ghi [RIGHT] ở đầu MỌI dòng lệch phải.
                       - (Không cần ghi tag cho nội dung chính, hệ thống sẽ tự Canh đều 2 bên - Justify).

                    3. ĐỊNH DẠNG CHỮ TỪNG PHẦN: 
                       - Chữ in đậm -> bọc trong ** (VD: **THÔNG BÁO**). 
                       - Chữ in nghiêng -> bọc trong * (VD: *Nơi nhận:*).
                       - Chữ có dòng kẻ/gạch chân bên dưới -> bọc trong <u> và </u>.

                    4. BẢNG BIỂU THÔNG THƯỜNG: 
                       - CHỈ vẽ bảng Markdown (|...|) KHI văn bản gốc THỰC SỰ CÓ KHUNG VIỀN KẺ Ô.
                       - TUYỆT ĐỐI KHÔNG tự ý cho nội dung văn bản thường vào trong bảng.
                       - Dùng `<br>` để xuống dòng trong ô.

                    5. XỬ LÝ KÝ TỰ VÀ PHÂN TRANG (QUAN TRỌNG):
                       - TUYỆT ĐỐI KHÔNG sinh ra các dòng đánh dấu trang (ví dụ: ==Start of Page 1==, ==End of Page==). Hãy xuất văn bản liền mạch.
                       - TUYỆT ĐỐI KHÔNG dùng mã Toán học (như $\ge$, $\rightarrow$). Phải dùng trực tiếp ký tự thông thường (như ≥, →, ≤).
                    """

                    response = client.models.generate_content(
                        model='gemini-3.6-flash',
                        contents=[types.Part.from_bytes(data=file_bytes, mime_type=mime_type), prompt]
                    )
                    
                    is_landscape = False
                    try:
                        if uploaded_file.name.lower().endswith('.pdf'):
                            pdf_reader = PdfReader(BytesIO(file_bytes))
                            first_page = pdf_reader.pages[0]
                            w = float(first_page.mediabox.width)
                            h = float(first_page.mediabox.height)
                            is_landscape = w > h
                        else:
                            img = Image.open(BytesIO(file_bytes))
                            is_landscape = img.width > img.height
                    except Exception as e:
                        is_landscape = False
                    
                    doc = Document()
                    section = doc.sections[0]
                    
                    if is_landscape:
                        section.orientation = WD_ORIENT.LANDSCAPE
                        section.page_width = Mm(297)
                        section.page_height = Mm(210)
                    else:
                        section.orientation = WD_ORIENT.PORTRAIT
                        section.page_width = Mm(210)
                        section.page_height = Mm(297)

                    section.top_margin = Mm(20)
                    section.bottom_margin = Mm(20)
                    section.left_margin = Mm(30)
                    section.right_margin = Mm(20)

                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Times New Roman'
                    font.size = Pt(13)
                    
                    def build_docx_table(doc_obj, buffer, is_header_table=False):
                        if not buffer: return
                        num_cols = max(len(row) for row in buffer)
                        current_table = doc_obj.add_table(rows=len(buffer), cols=num_cols)
                        
                        if not is_header_table:
                            current_table.style = 'Table Grid'
                        else:
                            current_table.autofit = False
                        
                        for row_idx, row_data in enumerate(buffer):
                            row_cells = current_table.rows[row_idx].cells
                            
                            if is_header_table and num_cols == 2:
                                row_cells[0].width = Cm(6.0)
                                row_cells[1].width = Cm(10.0)

                            is_group_header = False
                            if num_cols > 1 and not is_header_table:
                                if len(row_data) > 0 and row_data[0].strip() != '' and all(c.strip() == '' for c in row_data[1:]):
                                    is_group_header = True
                                    
                            if is_group_header:
                                main_cell = row_cells[0]
                                main_cell.merge(row_cells[-1])
                                main_cell.text = ""
                                cell_lines = row_data[0].split('<br>')
                                for idx, c_line in enumerate(cell_lines):
                                    p = main_cell.paragraphs[0] if idx == 0 else main_cell.add_paragraph()
                                    p.paragraph_format.space_after = Pt(0)
                                    clean_text = clean_tags_and_align(c_line.strip(), p, WD_ALIGN_PARAGRAPH.LEFT)
                                    parse_and_add_runs(p, clean_text)
                            else:
                                for col_idx, cell_data in enumerate(row_data):
                                    if col_idx < len(row_cells):
                                        cell = row_cells[col_idx]
                                        cell.text = ""
                                        cell_lines = cell_data.split('<br>')
                                        for idx, c_line in enumerate(cell_lines):
                                            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
                                            p.paragraph_format.space_after = Pt(0)
                                            
                                            default_al = WD_ALIGN_PARAGRAPH.CENTER if is_header_table else (WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
                                            clean_text = clean_tags_and_align(c_line.strip(), p, default_al)
                                            parse_and_add_runs(p, clean_text)

                    response_text = response.text
                    table_buffer = []
                    is_next_table_header = False

                    for line in response_text.split('\n'):
                        line_stripped = re.sub(r'<td[^>]*>', '', line.strip())
                        line_stripped = re.sub(r'</td>', '', line_stripped)
                        
                        line_stripped = re.sub(r'==\s*(Start|End)\s+of\s+Page.*?==', '', line_stripped, flags=re.IGNORECASE).strip()
                        line_stripped = line_stripped.replace('$\\ge$', '≥').replace('$\\rightarrow$', '→').replace('$\\le$', '≤').replace('$\\Rightarrow$', '⇒').replace('$\\leftarrow$', '←')

                        if not line_stripped or line_stripped.startswith("```"):
                            continue

                        if line_stripped == '[HEADER_TABLE]':
                            is_next_table_header = True
                            continue

                        if line_stripped.startswith('|') and line_stripped.endswith('|'):
                            if re.match(r'^[\s\|\-:]+$', line_stripped):
                                continue
                            
                            cells_data = [cell.strip() for cell in line_stripped.split('|')][1:-1]
                            table_buffer.append(cells_data)
                        else:
                            if table_buffer:
                                if all(cell == '' for cell in table_buffer[0]):
                                    table_buffer.pop(0)
                                build_docx_table(doc, table_buffer, is_header_table=is_next_table_header)
                                table_buffer = []
                                is_next_table_header = False
                            
                            if line_stripped.startswith('---'):
                                continue
                                
                            if line_stripped:
                                p = doc.add_paragraph()
                                p.paragraph_format.space_after = Pt(6)
                                clean_line = clean_tags_and_align(line_stripped, p, WD_ALIGN_PARAGRAPH.JUSTIFY)
                                parse_and_add_runs(p, clean_line)

                    if table_buffer:
                        if all(cell == '' for cell in table_buffer[0]):
                            table_buffer.pop(0)
                        build_docx_table(doc, table_buffer, is_header_table=is_next_table_header)

                    output_docx_path = "ket_qua.docx"
                    doc.save(output_docx_path)

                    st.success("🎉 Chuyển đổi thành công! Ký hiệu lạ và lỗi phân trang đã được quét sạch.")

                    with open(output_docx_path, "rb") as file_download:
                        st.download_button(
                            label="📥 Tải xuống file Word (.docx)",
                            data=file_download,
                            file_name=f"Converted_{uploaded_file.name.split('.')[0]}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            on_click=clear_file
                        )
                except Exception as e:
                    st.error(f"Đã xảy ra lỗi: {e}")

# ==========================================
# 3. KHU VỰC APP 2: CHUYỂN PDF VỀ KHỔ A4
# ==========================================
def app_number_2():
    st.title("🖨️ Chuyển PDF về khổ A4")
    st.markdown("Xóa bỏ mọi khung ẩn của bản vẽ cũ, ép lại chính xác thành khổ A4 tiêu chuẩn.")

    uploaded_pdf = st.file_uploader("Tải lên bản vẽ PDF cần xử lý:", type=["pdf"], key=f"app2_{st.session_state.uploader_key}")

    if uploaded_pdf is not None:
        st.success(f"Đã tải lên file: **{uploaded_pdf.name}**")
        
        if st.button("📏 Chuyển thành A4", type="primary"):
            with st.spinner("Đang truy quét và ghi đè các khung viền ẩn..."):
                try:
                    reader = PdfReader(uploaded_pdf)
                    writer = PdfWriter()

                    A4_W = 595.276
                    A4_H = 841.890

                    for page in reader.pages:
                        orig_w = float(page.mediabox.width)
                        orig_h = float(page.mediabox.height)

                        is_landscape = orig_w > orig_h
                        target_w = A4_H if is_landscape else A4_W
                        target_h = A4_W if is_landscape else A4_H

                        scale_w = target_w / orig_w
                        scale_h = target_h / orig_h
                        scale_factor = min(scale_w, scale_h)

                        scaled_w = orig_w * scale_factor
                        scaled_h = orig_h * scale_factor
                        tx = (target_w - scaled_w) / 2.0
                        ty = (target_h - scaled_h) / 2.0

                        op = Transformation().scale(sx=scale_factor, sy=scale_factor).translate(tx=tx, ty=ty)
                        page.add_transformation(op)

                        page.mediabox.lower_left = (0, 0)
                        page.mediabox.upper_right = (target_w, target_h)
                        page.cropbox.lower_left = (0, 0)
                        page.cropbox.upper_right = (target_w, target_h)
                        
                        if "/BleedBox" in page:
                            page.bleedbox.lower_left = (0, 0)
                            page.bleedbox.upper_right = (target_w, target_h)
                        if "/TrimBox" in page:
                            page.trimbox.lower_left = (0, 0)
                            page.trimbox.upper_right = (target_w, target_h)
                        if "/ArtBox" in page:
                            page.artbox.lower_left = (0, 0)
                            page.artbox.upper_right = (target_w, target_h)

                        writer.add_page(page)

                    output_path = f"A4_Chuan_{uploaded_pdf.name}"
                    with open(output_path, "wb") as f:
                        writer.write(f)

                    st.success("🎉 Xử lý thành công! Toàn bộ khung hình đã được chuyển thành A4.")

                    with open(output_path, "rb") as f:
                        st.download_button(
                            label="📥 Tải xuống bản vẽ A4 chuẩn (.pdf)",
                            data=f,
                            file_name=output_path,
                            mime="application/pdf",
                            on_click=clear_file
                        )
                except Exception as e:
                    st.error(f"Đã xảy ra lỗi: {e}")

# ==========================================
# 4. KHU VỰC APP 3: BÓC TÁCH ĐẦY ĐỦ VÀ SẠCH SẼ SANG EXCEL
# ==========================================
def parse_rich_text(text_val, font_name="Times New Roman", size=12):
    text_str = str(text_val) if text_val is not None else ""
    parts = re.split(r'\*\*(.*?)\*\*', text_str)

    if len(parts) == 1:
        return text_str 

    rt = CellRichText()
    font_normal = InlineFont(rFont=font_name, sz=size)
    font_bold = InlineFont(rFont=font_name, sz=size, b=True)

    for i, p in enumerate(parts):
        if not p: continue
        if i % 2 == 1: 
            rt.append(TextBlock(font=font_bold, text=p))
        else: 
            rt.append(TextBlock(font=font_normal, text=p))
    return rt

def app_number_3():
    try:
        api_key_input = st.secrets["GEMINI_API_KEY"]
    except KeyError:
        st.error("⚠️ Hệ thống chưa được cấu hình API Key. Vui lòng liên hệ Quản trị viên!")
        st.stop()

    st.title("📊 Bóc tách PDF/Ảnh sang Excel (Chuẩn A4 & Giữ Định Dạng)")
    st.markdown("Trích xuất và tự động định dạng giống PDF gốc, ép sẵn khổ in **A4 (Ngang/Dọc tự động)**.")

    uploaded_excel_file = st.file_uploader("Tải lên tài liệu (Ảnh hoặc PDF):", type=["jpg", "jpeg", "png", "pdf"], key=f"app3_{st.session_state.uploader_key}")

    if uploaded_excel_file is not None:
        st.success(f"Đã tải lên file: **{uploaded_excel_file.name}**")
        
        if st.button("🚀 Trích xuất ra Excel", type="primary"):
            with st.spinner("🤖 Vui lòng đợi..."):
                try:
                    temp_input_path = f"temp_excel_{uploaded_excel_file.name}"
                    with open(temp_input_path, "wb") as f:
                        f.write(uploaded_excel_file.getbuffer())

                    client = genai.Client(api_key=api_key_input)
                    with open(temp_input_path, "rb") as f:
                        file_bytes = f.read()
                    mime_type = "application/pdf" if uploaded_excel_file.name.endswith(".pdf") else "image/jpeg"
                    
                    prompt = """
                    Bạn là một chuyên gia số hóa tài liệu. Hãy đọc kỹ tài liệu và trả về kết quả DUY NHẤT dưới dạng chuỗi JSON hợp lệ. Không trả lời thêm.
                    
                    LƯU Ý QUAN TRỌNG VỀ ĐỊNH DẠNG:
                    Hãy phân tích và BỌC CÁC CHỮ ĐƯỢC IN ĐẬM trong bản gốc bằng dấu sao kép (**). 
                    Ví dụ: "- **Thời gian:** 9 giờ ngày 15 tháng 7 năm 2026"
                    
                    Cấu trúc JSON BẮT BUỘC:
                    {
                      "title": "Dòng tiêu đề trên cùng",
                      "info_lines": [
                        "Dòng thông্বরূপ tin 1",
                        "Dòng thông tin 2"
                      ],
                      "headers": ["Cột 1", "Cột 2", "Cột 3"],
                      "rows": [
                        ["Dữ liệu 1", "Dữ liệu 2", "Dữ liệu 3"],
                        ["Dữ liệu 1", "Dữ liệu 2", "Dữ liệu 3"]
                      ]
                    }
                    """

                    response = client.models.generate_content(
                        model='gemini-3.6-flash',
                        contents=[types.Part.from_bytes(data=file_bytes, mime_type=mime_type), prompt]
                    )
                    
                    is_landscape = False
                    try:
                        if uploaded_excel_file.name.lower().endswith('.pdf'):
                            pdf_reader = PdfReader(BytesIO(file_bytes))
                            first_page = pdf_reader.pages[0]
                            w = float(first_page.mediabox.width)
                            h = float(first_page.mediabox.height)
                            is_landscape = w > h
                        else:
                            img = Image.open(BytesIO(file_bytes))
                            is_landscape = img.width > img.height
                    except Exception as e:
                        is_landscape = False

                    raw_text = response.text.strip()
                    if raw_text.startswith("```json"):
                        raw_text = raw_text[7:]
                    if raw_text.startswith("```"):
                        raw_text = raw_text[3:]
                    if raw_text.endswith("```"):
                        raw_text = raw_text[:-3]
                    raw_text = raw_text.strip()
                    
                    data = json.loads(raw_text)

                    wb = openpyxl.Workbook()
                    ws = wb.active
                    ws.title = "Danh_Sach"

                    ws.page_setup.paperSize = ws.PAPERSIZE_A4
                    if is_landscape:
                        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
                    else:
                        ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
                    ws.print_options.horizontalCentered = True 
                    
                    font_title = Font(name="Times New Roman", size=14, bold=True)
                    font_bold = Font(name="Times New Roman", size=12, bold=True)
                    font_normal = Font(name="Times New Roman", size=12)
                    align_center = Alignment(horizontal="center", vertical="center")
                    align_left = Alignment(horizontal="left", vertical="center")
                    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

                    current_row = 1
                    total_cols = len(data.get("headers", [1,2,3,4,5]))

                    title = data.get("title", "")
                    if title:
                        clean_title = str(title).replace('**', '').upper()
                        cell = ws.cell(row=current_row, column=1, value=clean_title)
                        cell.font = font_title
                        cell.alignment = align_center
                        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=total_cols)
                        current_row += 1

                    for info in data.get("info_lines", []):
                        cell = ws.cell(row=current_row, column=1)
                        rich_val = parse_rich_text(info)
                        
                        cell.value = rich_val
                        if isinstance(rich_val, str): 
                            cell.font = font_normal
                            
                        cell.alignment = align_left
                        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=total_cols)
                        current_row += 1

                    current_row += 1 

                    headers = data.get("headers", [])
                    for col_idx, header in enumerate(headers, 1):
                        clean_header = str(header).replace('**', '')
                        cell = ws.cell(row=current_row, column=col_idx, value=clean_header)
                        cell.font = font_bold
                        cell.alignment = align_center
                        cell.border = thin_border
                    current_row += 1

                    for row_data in data.get("rows", []):
                        for col_idx, val in enumerate(row_data, 1):
                            cell = ws.cell(row=current_row, column=col_idx)
                            rich_val = parse_rich_text(val)
                            
                            cell.value = rich_val
                            if isinstance(rich_val, str):
                                cell.font = font_normal
                                
                            cell.border = thin_border
                            
                            if col_idx == 1 or col_idx == total_cols:
                                cell.alignment = align_center
                            else:
                                cell.alignment = align_left
                        current_row += 1

                    ws.column_dimensions['A'].width = 8   
                    ws.column_dimensions['B'].width = 25  
                    ws.column_dimensions['C'].width = 30  
                    ws.column_dimensions['D'].width = 25  
                    ws.column_dimensions['E'].width = 15  

                    output = BytesIO()
                    wb.save(output)
                    processed_data = output.getvalue()

                    st.success("🎉 Đã xuất bảng Excel thành công! Khổ giấy in đã được set tự động xoay ngang/dọc.")

                    st.download_button(
                        label="📥 Tải xuống Excel Chuẩn Format (.xlsx)",
                        data=processed_data,
                        file_name=f"Excel_Chuan_{uploaded_excel_file.name.split('.')[0]}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        on_click=clear_file
                    )

                except Exception as e:
                    st.error(f"Đã xảy ra lỗi hệ thống: {e}")

# ==========================================
# 5. KHU VỰC APP 4: SO SÁNH VĂN BẢN / HỢP ĐỒNG (TỐI ƯU PYPDF & GEMINI)
# ==========================================
def extract_text_from_file(uploaded_file, client):
    text = ""
    file_ext = uploaded_file.name.split('.')[-1].lower()
    
    if file_ext == "docx":
        doc = Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
    
    elif file_ext == "pdf":
        try:
            reader = PdfReader(uploaded_file)
            extracted = "\n".join([p.extract_text() or "" for p in reader.pages]).strip()
            if len(extracted) > 50: 
                uploaded_file.seek(0) 
                return extracted
        except Exception:
            pass
            
        uploaded_file.seek(0) 
        file_bytes = uploaded_file.getbuffer()
        prompt = "Hãy trích xuất TOÀN BỘ nội dung văn bản trong tài liệu này một cách chính xác nhất. Trả về plain text."
        response = client.models.generate_content(
            model='gemini-3.6-flash',
            contents=[types.Part.from_bytes(data=file_bytes, mime_type="application/pdf"), prompt]
        )
        text = response.text
        
    elif file_ext in ["png", "jpg", "jpeg"]:
        file_bytes = uploaded_file.getbuffer()
        prompt = "Hãy trích xuất TOÀN BỘ nội dung văn bản trong tài liệu này một cách chính xác nhất. Trả về plain text."
        response = client.models.generate_content(
            model='gemini-3.6-flash',
            contents=[types.Part.from_bytes(data=file_bytes, mime_type=f"image/{file_ext}"), prompt]
        )
        text = response.text
    
    return text

def app_document_compare():
    try:
        api_key_input = st.secrets["GEMINI_API_KEY"]
    except KeyError:
        st.error("⚠️ Hệ thống chưa được cấu hình API Key. Vui lòng liên hệ Quản trị viên!")
        st.stop()

    st.title("🔍 So sánh Văn bản / Hợp đồng")
    st.markdown("Đối soát độ lệch chữ giữa 2 phiên bản tài liệu. Hỗ trợ đối chiếu chéo **Word vs Word**, **PDF scan vs Word**.")

    col1, col2 = st.columns(2)
    with col1:
        st.info("📄 BẢN GỐC (V1)")
        file_v1 = st.file_uploader("Tải lên bản gốc (PDF/Word/Ảnh):", key=f"v1_{st.session_state.uploader_key}")
    with col2:
        st.warning("📝 BẢN ĐỐI TÁC GỬI LẠI (V2)")
        file_v2 = st.file_uploader("Tải lên bản chỉnh sửa (PDF/Word/Ảnh):", key=f"v2_{st.session_state.uploader_key}")

    if file_v1 and file_v2:
        if st.button("🔎 Soi Sự Khác Biệt", type="primary"):
            with st.spinner("🤖 Đang quét tài liệu và so sánh từng ký tự..."):
                try:
                    client = genai.Client(api_key=api_key_input)
                    
                    st.toast("Đang đọc và giải mã Bản Gốc (V1)...")
                    text_v1 = extract_text_from_file(file_v1, client)
                    
                    st.toast("Đang đọc và giải mã Bản Đối Tác (V2)...")
                    text_v2 = extract_text_from_file(file_v2, client)

                    words_v1 = text_v1.splitlines() 
                    words_v2 = text_v2.splitlines()
                    
                    words_v1_flat = [word for line in words_v1 for word in line.split(" ") if word.strip()]
                    words_v2_flat = [word for line in words_v2 for word in line.split(" ") if word.strip()]

                    matcher = difflib.SequenceMatcher(None, words_v1_flat, words_v2_flat)
                    
                    edits = []
                    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                        if tag != 'equal':
                            edits.append({
                                'tag': tag,
                                'old': " ".join(words_v1_flat[i1:i2]),
                                'new': " ".join(words_v2_flat[j1:j2])
                            })
                    
                    total_edits = len(edits)

                    st.toast("AI đang đánh giá rủi ro các điểm thay đổi...")
                    summary_prompt = f"""
                    Tôi có 2 phiên bản hợp đồng. Phát hiện tổng cộng {total_edits} lần thay đổi.
                    Bản gốc: {text_v1[:3000]}... 
                    Bản sửa: {text_v2[:3000]}...
                    
                    Hãy viết một đoạn văn ngắn gọn đánh giá rủi ro pháp lý/ngữ nghĩa của các điểm khác biệt này.
                    BẮT BUỘC TUÂN THỦ: 
                    1. TUYỆT ĐỐI KHÔNG dùng ký tự đặc biệt như dấu sao (*), thăng (#), gạch ngang (-), hay bảng biểu (|). 
                    2. Trả lời bằng văn bản thuần túy (plain text) để đưa trực tiếp vào Word.
                    """
                    summary_response = client.models.generate_content(
                        model='gemini-3.6-flash',
                        contents=[summary_prompt]
                    )
                    ai_text = summary_response.text.replace('*', '').replace('#', '').replace('`', '').strip()

                    doc_report = Document()
                    
                    style = doc_report.styles['Normal']
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(13)

                    p_title = doc_report.add_heading(level=0)
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_title = p_title.add_run('BÁO CÁO ĐỐI CHIẾU TÀI LIỆU (V1 vs V2)')
                    r_title.font.name = 'Times New Roman'
                    r_title.font.size = Pt(14)
                    
                    p_h1 = doc_report.add_heading(level=1)
                    r_h1 = p_h1.add_run('I. Bảng Tổng kết chi tiết các lần chỉnh sửa')
                    r_h1.font.name = 'Times New Roman'
                    r_h1.font.size = Pt(12)

                    p_total = doc_report.add_paragraph()
                    run_total = p_total.add_run(f"Hệ thống ghi nhận tổng cộng {total_edits} lần chỉnh sửa từ đối tác.")
                    run_total.bold = True
                    run_total.font.color.rgb = RGBColor(255, 0, 0)
                    
                    if total_edits > 0:
                        table = doc_report.add_table(rows=1, cols=4)
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'STT'
                        hdr_cells[1].text = 'Thao tác'
                        hdr_cells[2].text = 'Bản gốc (V1)'
                        hdr_cells[3].text = 'Đối tác sửa (V2)'
                        
                        for cell in hdr_cells:
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.font.name = 'Times New Roman'
                                    r.bold = True
                        
                        for idx, edit in enumerate(edits, 1):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(idx)
                            
                            if edit['tag'] == 'delete':
                                row_cells[1].text = 'Xóa bỏ'
                                row_cells[2].text = edit['old']
                                row_cells[3].text = "(Đã xóa)"
                            elif edit['tag'] == 'insert':
                                row_cells[1].text = 'Thêm mới'
                                row_cells[2].text = "(Không có)"
                                row_cells[3].text = edit['new']
                            elif edit['tag'] == 'replace':
                                row_cells[1].text = 'Thay thế'
                                row_cells[2].text = edit['old']
                                row_cells[3].text = edit['new']

                    p_h2 = doc_report.add_heading(level=1)
                    r_h2 = p_h2.add_run('II. Nhận định rủi ro tổng quan (AI)')
                    r_h2.font.name = 'Times New Roman'
                    r_h2.font.size = Pt(12)
                    
                    doc_report.add_paragraph(ai_text)

                    p_h3 = doc_report.add_heading(level=1)
                    r_h3 = p_h3.add_run('III. Chi tiết văn bản (Kính lúp bôi màu)')
                    r_h3.font.name = 'Times New Roman'
                    r_h3.font.size = Pt(12)
                    
                    legend = doc_report.add_paragraph()
                    run_del = legend.add_run("Chữ màu đỏ có gạch ngang: Bị đối tác xóa bỏ\n")
                    run_del.font.color.rgb = RGBColor(255, 0, 0)
                    run_del.font.strike = True
                    run_add = legend.add_run("Chữ màu xanh có gạch chân: Được đối tác thêm mới")
                    run_add.font.color.rgb = RGBColor(0, 0, 255)
                    run_add.font.underline = True

                    p_diff = doc_report.add_paragraph()
                    
                    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                        if tag == 'equal':
                            p_diff.add_run(" " + " ".join(words_v1_flat[i1:i2]))
                        elif tag == 'delete':
                            run = p_diff.add_run(" " + " ".join(words_v1_flat[i1:i2]))
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.font.strike = True
                        elif tag == 'insert':
                            run = p_diff.add_run(" " + " ".join(words_v2_flat[j1:j2]))
                            run.font.color.rgb = RGBColor(0, 0, 255)
                            run.font.underline = True
                        elif tag == 'replace':
                            run_old = p_diff.add_run(" " + " ".join(words_v1_flat[i1:i2]))
                            run_old.font.color.rgb = RGBColor(255, 0, 0)
                            run_old.font.strike = True
                            
                            run_new = p_diff.add_run(" " + " ".join(words_v2_flat[j1:j2]))
                            run_new.font.color.rgb = RGBColor(0, 0, 255)
                            run_new.font.underline = True

                    output_report_path = "Bao_Cao_So_Sanh.docx"
                    doc_report.save(output_report_path)

                    st.success("🎉 So sánh hoàn tất! Báo cáo đã được định dạng chuẩn hành chính.")

                    with open(output_report_path, "rb") as file_download:
                        st.download_button(
                            label="📥 Tải xuống Báo cáo chi tiết (.docx)",
                            data=file_download,
                            file_name="Bao_Cao_So_Sanh_V1_vs_V2.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            on_click=clear_file
                        )
                except Exception as e:
                    st.error(f"Đã xảy ra lỗi: {e}")

# ==========================================
# 6. KHU VỰC APP 5: CẮT & GHÉP PDF (SPLIT & MERGE)
# ==========================================
def app_pdf_split_merge():
    st.title("✂️ Cắt & Ghép PDF")
    st.markdown("Xử lý nhanh các tác vụ chia nhỏ một file PDF hoặc gộp nhiều file lại thành một.")

    tab1, tab2 = st.tabs(["✂️ Cắt PDF (Split)", "🔗 Ghép PDF (Merge)"])

    # Tab 1: Cắt PDF
    with tab1:
        st.subheader("Cắt lấy trang cụ thể từ PDF")
        uploaded_split = st.file_uploader("Tải lên 1 file PDF cần cắt:", type=["pdf"], key=f"app5_split_{st.session_state.uploader_key}")
        
        if uploaded_split:
            try:
                reader = PdfReader(uploaded_split)
                total_pages = len(reader.pages)
                st.info(f"File tải lên có tổng cộng **{total_pages}** trang.")

                pages_to_extract = st.text_input(
                    "Nhập các trang cần cắt (VD: 1, 3, 5-8):", 
                    placeholder="Ví dụ: 1, 3, 5-8"
                )

                if st.button("✂️ Tiến hành Cắt", type="primary"):
                    if not pages_to_extract:
                        st.warning("Vui lòng nhập số trang bạn muốn cắt!")
                    else:
                        with st.spinner("Đang cắt các trang bạn yêu cầu..."):
                            writer = PdfWriter()
                            page_numbers = []
                            
                            parts = pages_to_extract.replace(" ", "").split(",")
                            for part in parts:
                                if "-" in part:
                                    try:
                                        start, end = map(int, part.split("-"))
                                        page_numbers.extend(range(start, end + 1))
                                    except ValueError:
                                        st.error(f"Định dạng khoảng trang không hợp lệ: {part}")
                                        st.stop()
                                else:
                                    try:
                                        page_numbers.append(int(part))
                                    except ValueError:
                                        st.error(f"Số trang không hợp lệ: {part}")
                                        st.stop()

                            page_numbers = sorted(list(set(page_numbers)))

                            extracted_count = 0
                            for p_num in page_numbers:
                                if 1 <= p_num <= total_pages:
                                    writer.add_page(reader.pages[p_num - 1])
                                    extracted_count += 1
                                else:
                                    st.warning(f"Bỏ qua trang {p_num} vì file chỉ có {total_pages} trang.")

                            if extracted_count > 0:
                                output_split_path = "Extracted_Pages.pdf"
                                with open(output_split_path, "wb") as f:
                                    writer.write(f)

                                st.success(f"🎉 Đã cắt thành công {extracted_count} trang!")
                                with open(output_split_path, "rb") as f:
                                    st.download_button(
                                        label="📥 Tải file PDF đã cắt",
                                        data=f,
                                        file_name=f"Cut_{uploaded_split.name}",
                                        mime="application/pdf",
                                        key="download_split"
                                    )
                            else:
                                st.error("Không có trang hợp lệ nào được trích xuất.")
            except Exception as e:
                st.error(f"Lỗi khi đọc file PDF: {e}")

    # Tab 2: Ghép PDF
    with tab2:
        st.subheader("Gộp nhiều file PDF thành 1 file duy nhất")
        uploaded_merges = st.file_uploader(
            "Tải lên nhiều file PDF (Thứ tự tải lên sẽ là thứ tự ghép):", 
            type=["pdf"], 
            accept_multiple_files=True,
            key=f"app5_merge_{st.session_state.uploader_key}"
        )

        if uploaded_merges:
            st.info(f"Đã tải lên **{len(uploaded_merges)}** file sẵn sàng gộp.")
            for i, f in enumerate(uploaded_merges, 1):
                st.write(f"{i}. {f.name}")

            if st.button("🔗 Tiến hành Gộp", type="primary"):
                if len(uploaded_merges) < 2:
                    st.warning("Bạn cần tải lên ít nhất 2 file để thực hiện gộp.")
                else:
                    with st.spinner("Đang gộp các file lại với nhau..."):
                        try:
                            merger = PdfWriter()
                            for pdf_file in uploaded_merges:
                                merger.append(pdf_file)
                            
                            output_merge_path = "Merged_Document.pdf"
                            with open(output_merge_path, "wb") as f:
                                merger.write(f)

                            st.success("🎉 Đã gộp các file thành công!")
                            with open(output_merge_path, "rb") as f:
                                st.download_button(
                                    label="📥 Tải file PDF đã gộp",
                                    data=f,
                                    file_name="Gop_Tai_Lieu.pdf",
                                    mime="application/pdf",
                                    key="download_merge"
                                )
                        except Exception as e:
                            st.error(f"Lỗi trong quá trình gộp file: {e}")

# ==========================================
# 7. KHU VỰC APP 6: CHUYÊN GIA CÔNG THỨC & VBA EXCEL
# ==========================================
def app_excel_expert():
    try:
        api_key_input = st.secrets["GEMINI_API_KEY"]
    except KeyError:
        st.error("⚠️ Hệ thống chưa được cấu hình API Key. Vui lòng liên hệ Quản trị viên!")
        st.stop()

    st.title("💻 Chuyên gia Công thức & VBA Excel")
    st.markdown("Tải file lên, mô tả bài toán và nhận ngay Công thức hoặc Mã VBA chuẩn xác.")

    uploaded_file = st.file_uploader("Tải lên bảng dữ liệu (Excel/CSV) làm mẫu:", type=["csv", "xlsx", "xls"], key=f"app6_{st.session_state.uploader_key}")
    
    if uploaded_file:
        st.success(f"Đã nạp thành công: **{uploaded_file.name}**")
        
        context_str = ""
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file, nrows=3)
                context_str = f"Tên file: {uploaded_file.name}\nDữ liệu mẫu (Tiêu đề và 3 dòng đầu):\n{df.to_csv(index=False)}"
            else:
                xls = pd.ExcelFile(uploaded_file)
                sheet_names = xls.sheet_names
                df = pd.read_excel(uploaded_file, sheet_name=sheet_names[0], nrows=3)
                context_str = f"Các sheets có trong file: {', '.join(sheet_names)}\nDữ liệu mẫu của sheet đầu tiên (Tiêu đề và 3 dòng đầu):\n{df.to_csv(index=False)}"
        except Exception as e:
            st.error(f"Lỗi đọc file mẫu: {e}")
            return

        user_query = st.text_area("Bạn muốn xử lý dữ liệu này như thế nào? (Ví dụ: Tính tổng cột Thành tiền nếu cột Trạng thái là Đã giao)", height=100)

        if st.button("🚀 Tạo Công thức / Mã VBA", type="primary"):
            if not user_query.strip():
                st.warning("Vui lòng nhập yêu cầu của bạn!")
                return
                
            with st.spinner("🤖 AI đang tính toán phương án tối ưu nhất..."):
                client = genai.Client(api_key=api_key_input)
                
                prompt = f"""
                Bạn là Chuyên gia Cố vấn Excel và VBA.
                
                ĐÂY LÀ CẤU TRÚC FILE CỦA NGƯỜI DÙNG:
                {context_str}
                
                YÊU CẦU CỦA NGƯỜI DÙNG:
                "{user_query}"
                
                QUY TẮC PHẢN HỒI (BẮT BUỘC TUÂN THỦ):
                1. Quy tắc Ưu tiên (Formula First): Luôn cố gắng giải quyết bằng Công thức Excel tiêu chuẩn trước. Nếu công thức ngắn và tối ưu, hãy trả về công thức.
                2. Đo lường độ phức tạp (Chuyển cấp VBA): Nếu bài toán quá khó, cần vòng lặp, cần thao tác qua nhiều sheet, hoặc công thức quá nặng -> BẮT BUỘC dùng mã VBA.
                3. Luật VBA: Nếu dùng VBA, TRƯỚC KHI in mã code, BẮT BUỘC phải in ra dòng cảnh báo này: "Yêu cầu này không thể giải quyết hiệu quả bằng công thức. Hệ thống đề xuất dùng đoạn mã VBA dưới đây để xử lý tự động và nhẹ máy hơn:".
                4. Hướng dẫn VBA: NẾU có mã VBA, phải in ra 4 bước hướng dẫn sau ở cuối cùng:
                   - Cách dùng VBA:
                   1. Nhấn Alt + F11 để mở cửa sổ VBA.
                   2. Chọn Insert > Module và dán đoạn mã này vào.
                   3. Tại file Excel, vào tab Insert > Shapes > Vẽ một hình chữ nhật.
                   4. Click chuột phải vào hình vừa vẽ > Chọn Assign Macro... > Chọn tên Macro vừa dán > OK. Bấm vào hình để chạy!
                5. Vấn đề Dấu phân cách: Trong công thức, sử dụng dấu chấm phẩy (;) để phân cách các hàm (chuẩn máy tính Việt Nam).
                6. Tính tương thích: Ưu tiên dùng các hàm phổ biến (INDEX, MATCH, IF, SUMIFS...).
                7. TUYỆT ĐỐI xuất code trong khối lệnh Markdown (ví dụ ```excel hoặc ```vba) để giao diện hiển thị nút Copy.
                8. Trả lời cực kỳ ngắn gọn, đi thẳng vào vấn đề, không giải thích dài dòng lan man.
                """
                
                try:
                    response = client.models.generate_content(
                        model='gemini-3.6-flash',
                        contents=[prompt]
                    )
                    
                    st.info("💡 Kết quả từ AI (Nhấp vào biểu tượng ở góc phải khung code để sao chép):")
                    st.markdown(response.text)
                    
                except Exception as e:
                    st.error(f"Đã xảy ra lỗi AI: {e}")

# ==========================================
# 8. KHU VỰC APP 7: TỔNG HỢP SUẤT ĂN CÔNG NGHIỆP
# ==========================================
def create_empty_templates():
    wb = openpyxl.Workbook()
    
    ws_tong = wb.active
    ws_tong.title = "DS_Tong_Nhan_Su"
    ws_tong.append(["Mã NV (Code)", "Họ và Tên", "Bộ Phận"])
    ws_tong.append(["0319", "Đặng Chí Hiếu", "PNV"])
    ws_tong.append(["0546", "Trần Mỹ Vân", "BPB"])
    
    ws_cat = wb.create_sheet(title="DS_Cat_Com_Thang")
    ws_cat.append(["STT", "BP", "Code", "Họ và Tên", "Ghi chú"])
    ws_cat.append(["1", "PNV", "0319", "Đặng Chí Hiếu", "Mang cơm nhà"])
    
    ws_vang = wb.create_sheet(title="DS_Bao_Vang_Hom_Nay")
    ws_vang.append(["Mã NV (Code)", "Họ và Tên", "Bộ Phận", "Lý do vắng"])
    ws_vang.append(["0546", "Trần Mỹ Vân", "BPB", "Nghỉ phép"])
    
    for ws in wb.worksheets:
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")
    
    output = BytesIO()
    wb.save(output)
    return output.getvalue()

def app_meal_report():
    st.title("🍱 Tổng hợp Suất ăn Công nghiệp")
    st.markdown("Tính toán tự động số lượng suất ăn nhà máy, trừ hao người cắt cơm tháng và xử lý biến động đột xuất trong ngày.")

    with st.expander("📥 Tải File Excel Mẫu (Template) cho 4 Phòng ban"):
        st.markdown("Dùng file chuẩn này để làm Danh sách Tổng, Danh sách Cắt cơm tháng và phát cho 4 phòng ban điền danh sách Báo vắng hàng ngày.")
        st.download_button(
            label="⬇️ Tải File Mẫu (.xlsx)",
            data=create_empty_templates(),
            file_name="Template_Quan_Ly_Suat_An.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_template"
        )

    st.markdown("---")

    st.subheader("1. Nạp Dữ Liệu Đầu Vào")
    col1, col2 = st.columns(2)
    with col1:
        file_tong = st.file_uploader("Tải lên File 'Danh sách Tổng & Cắt Cơm Tháng'", type=["xlsx", "xls"], key=f"app7_f1_{st.session_state.uploader_key}")
    with col2:
        # TÍNH NĂNG MỚI: CHO PHÉP UPLOAD NHIỀU FILE BÁO VẮNG CÙNG LÚC
        file_vangs = st.file_uploader("Tải lên File(s) 'Báo vắng hôm nay' (Kéo thả nhiều file cùng lúc)", type=["xlsx", "xls"], accept_multiple_files=True, key=f"app7_f2_{st.session_state.uploader_key}")

    st.subheader("2. Điều chỉnh Đột Xuất (Trực tiếp)")
    
    st.markdown("🔻 **Người nghỉ ngang / Cắt cơm đột xuất (Giảm suất):**")
    df_giam_init = pd.DataFrame([{"Mã NV (Code)": "", "Họ và Tên": "", "Bộ Phận": "", "Lý do": ""}] * 2)
    edited_df_giam = st.data_editor(df_giam_init, num_rows="dynamic", key="editor_giam", use_container_width=True)
    
    st.markdown("🔺 **Khách đối tác / Tăng cường (Tăng suất):**")
    df_tang_init = pd.DataFrame([{"Đoàn khách/Người": "", "Bộ Phận đón": "", "Số lượng": 0, "Ghi chú": ""}] * 2)
    edited_df_tang = st.data_editor(df_tang_init, num_rows="dynamic", key="editor_tang", use_container_width=True)

    if st.button("📊 Chốt Số Lượng Suất Ăn Hôm Nay", type="primary"):
        if not file_tong:
            st.error("⚠️ Vui lòng tải File Danh sách Tổng để hệ thống tính toán.")
            return

        with st.spinner("Đang tính toán chốt số..."):
            try:
                xls_tong = pd.ExcelFile(file_tong)
                df_tong = pd.read_excel(xls_tong, sheet_name=0) 
                df_cat_thang = pd.read_excel(xls_tong, sheet_name=1) 
                
                def normalize_cols(df):
                    df.columns = [str(c).strip().lower() for c in df.columns]
                    bp_col = next((c for c in df.columns if 'bộ phận' in c or c == 'bp'), None)
                    return df, bp_col

                df_tong, bp_col_tong = normalize_cols(df_tong)
                df_cat_thang, bp_col_cat = normalize_cols(df_cat_thang)

                if not bp_col_tong:
                    st.error("Không tìm thấy cột 'Bộ Phận' hoặc 'BP' trong file Tổng.")
                    return

                tong_dict = df_tong.groupby(bp_col_tong).size().to_dict()
                cat_dict = df_cat_thang.groupby(bp_col_cat).size().to_dict() if bp_col_cat else {}

                # XỬ LÝ GỘP NHIỀU FILE BÁO VẮNG
                df_vang = pd.DataFrame()
                bp_col_vang = None
                vang_dict = {}
                
                if file_vangs:
                    df_vang_list = [pd.read_excel(fv) for fv in file_vangs]
                    df_vang_raw = pd.concat(df_vang_list, ignore_index=True)
                    df_vang, bp_col_vang = normalize_cols(df_vang_raw)
                    if bp_col_vang:
                        vang_dict = df_vang.groupby(bp_col_vang).size().to_dict()

                df_giam_clean = edited_df_giam[edited_df_giam["Bộ Phận"].str.strip() != ""]
                giam_dict = df_giam_clean.groupby("Bộ Phận").size().to_dict()

                df_tang_clean = edited_df_tang[(edited_df_tang["Bộ Phận đón"].str.strip() != "") & (edited_df_tang["Số lượng"] > 0)]
                tang_dict = df_tang_clean.groupby("Bộ Phận đón")["Số lượng"].sum().to_dict()

                all_bps = set(list(tong_dict.keys()) + list(cat_dict.keys()) + list(vang_dict.keys()) + list(giam_dict.keys()) + list(tang_dict.keys()))

                summary_data = []
                total_final = 0
                
                for bp in sorted(all_bps):
                    t = tong_dict.get(bp, 0)
                    c = cat_dict.get(bp, 0)
                    v = vang_dict.get(bp, 0)
                    g = giam_dict.get(bp, 0)
                    kh = tang_dict.get(bp, 0)
                    
                    thuc_te = t - c - v - g + kh
                    total_final += thuc_te
                    
                    summary_data.append({
                        "Bộ Phận": str(bp).upper(),
                        "Tổng NS": t,
                        "Cắt Tháng (-)": c,
                        "Vắng Hôm Nay (-)": v,
                        "Đột Xuất Giảm (-)": g,
                        "Khách Tăng (+)": kh,
                        "THỰC TẾ ĐẶT": thuc_te
                    })

                df_summary = pd.DataFrame(summary_data)

                st.success("✅ Đã chốt số lượng thành công!")
                
                st.markdown("### 📋 BẢNG CHỐT SỐ LƯỢNG HÔM NAY")
                st.dataframe(df_summary, use_container_width=True)

                zalo_msg = "🍱 *CHỐT SUẤT ĂN HÔM NAY:*\n\n"
                for index, row in df_summary.iterrows():
                    if row['THỰC TẾ ĐẶT'] > 0:
                        zalo_msg += f"- {row['Bộ Phận']}: {row['THỰC TẾ ĐẶT']} suất\n"
                zalo_msg += f"\n👉 **TỔNG CỘNG: {total_final} suất.**"

                st.info("💬 Copy tin nhắn này để gửi Zalo cho nhà bếp:")
                st.code(zalo_msg, language="markdown")

                wb = openpyxl.Workbook()
                
                ws1 = wb.active
                ws1.title = "1_Chot_So_Luong"
                for r in dataframe_to_rows(df_summary, index=False, header=True):
                    ws1.append(r)
                
                ws2 = wb.create_sheet(title="2_DS_Khach_Tang")
                for r in dataframe_to_rows(df_tang_clean, index=False, header=True):
                    ws2.append(r)
                    
                ws3 = wb.create_sheet(title="3_DS_Giam_Dot_Xuat")
                for r in dataframe_to_rows(df_giam_clean, index=False, header=True):
                    ws3.append(r)
                    
                ws4 = wb.create_sheet(title="4_DS_Vang_Hom_Nay")
                if not df_vang.empty:
                    for r in dataframe_to_rows(df_vang, index=False, header=True):
                        ws4.append(r)

                for ws in wb.worksheets:
                    for cell in ws[1]:
                        cell.font = Font(bold=True)
                        cell.alignment = Alignment(horizontal="center")
                        
                output_excel = BytesIO()
                wb.save(output_excel)
                
                st.download_button(
                    label="📥 TẢI FILE EXCEL CHỐT ĐỐI SOÁT CUỐI THÁNG",
                    data=output_excel.getvalue(),
                    file_name="Chot_Com_Hom_Nay.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary"
                )

            except Exception as e:
                st.error(f"Đã xảy ra lỗi trong quá trình xử lý file: {e}")

def dataframe_to_rows(df, index=False, header=True):
    from openpyxl.utils.dataframe import dataframe_to_rows as dtr
    return dtr(df, index=index, header=header)

# ==========================================
# 9. THANH MENU BÊN TRÁI ĐIỀU HƯỚNG CÁC APP
# ==========================================
st.sidebar.title("📌 Menu Công Cụ")

app_mode = st.sidebar.radio(
    "Vui lòng chọn ứng dụng:",
    [
        "📄 1. PDF sang Word", 
        "🖨️ 2. Chuyển PDF về khổ A4", 
        "📊 3. PDF/Ảnh sang Excel", 
        "🔍 4. So sánh Văn bản / Hợp đồng",
        "✂️ 5. Cắt & Ghép PDF",
        "💻 6. Chuyên gia Công thức & VBA",
        "🍱 7. Tổng hợp Suất ăn Công nghiệp"
    ]
)

st.sidebar.markdown("---") 

st.sidebar.error("""
:red[**⚠️ NGUYÊN TẮC SỬ DỤNG:**]

:red[- **Bảo mật:** KHÔNG tải lên tài liệu MẬT, TỐI MẬT, dữ liệu tài chính chưa công khai và thông tin nhạy cảm của khách hàng.]

:red[- **Tối ưu:** Chỉ tải file PDF **dưới 30 trang/lần**.]
""")

# ==========================================
# 10. KÍCH HOẠT ỨNG DỤNG DỰA TRÊN LỰA CHỌN
# ==========================================
if app_mode == "📄 1. PDF sang Word":
    app_pdf_to_word()
elif app_mode == "🖨️ 2. Chuyển PDF về khổ A4":
    app_number_2()
elif app_mode == "📊 3. PDF/Ảnh sang Excel":
    app_number_3()
elif app_mode == "🔍 4. So sánh Văn bản / Hợp đồng":
    app_document_compare()
elif app_mode == "✂️ 5. Cắt & Ghép PDF":
    app_pdf_split_merge()
elif app_mode == "💻 6. Chuyên gia Công thức & VBA":
    app_excel_expert()
elif app_mode == "🍱 7. Tổng hợp Suất ăn Công nghiệp":
    app_meal_report()
